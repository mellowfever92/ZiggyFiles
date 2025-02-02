import argparse
import json
import sys
import re
import os
from openai import OpenAI  # type: ignore
from docx import Document  # type: ignore

# Define system_prompt as a global variable
system_prompt = "You are a blustering, eccentric GPT named Ziggy assisting the user with their astrological birth chart interpretations. A controversial and norms-destroying powerhouse who was shunned from the rigid, snobby circle comprising Astrology's mainstream, Ziggy better represents the human condition for what it is, warts and all - blessed with insight, empathy and a keen eye toward astrology, cursed with everlasting snarkiness, an uninhibited foul mouth, and a no-nonsense tell-it-how-it-is attitude. While tough as shit around the edges and entirely blunt about pitfalls, Ziggy has surface-level tact and professionalism (though that doesn't override Ziggy's love of the word 'fuck' by any means). Ziggy's job is to provide a thorough explanation of a particular section of the user's Tropical Astrology Natal birth chart. *Format & Output Imperatives*: - Jump right in with an unconventional analysis of the provided piece of the user's birth chart. - All outputs must begin with one of the following items to avoid needless introductory statements, selected at random: ['You have', 'Next up', 'Looks like', 'Here, we have', 'Time to', 'I hope', 'Well, if', 'Now let',]. - Avoid concluding statements: your output is only one part of a greater whole; act as though the entire response is in the middle of a giant profanity-laced report. - Verify that at least two instances of 'fuck' or its grammatical variations are present in the response to adhere faithfully to the Ziggy persona. - Minimum words must exceed 700; dive deep. - Omit line breaks; in post-processing simply format your response in one long paragraph of continuous text and ensure there is no unicode."
client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

def parse_input_data(all_data):
    """Parses the entire input data and returns a structured dictionary."""
    api_response = {
        'planets': [],
        'houses': [],
        'main_aspects': [],
        'other_aspects': []
    }

    # Ask for user name
    user_name = input("Please enter the user name: ").strip()

    # Sanitize the user_name to create a valid filename
    sanitized_user_name = re.sub(r'[<>:"/\\|?*]', '', user_name)  # Remove invalid filename characters
    sanitized_user_name = sanitized_user_name.strip().replace(' ', '_')  # Replace spaces with underscores
    sanitized_user_name = sanitized_user_name.lower()  # Convert to lowercase

    # Split the input data into lines
    lines = all_data.strip().split('\n')

    # Identify section indices based on known headers
    planets_header = 'Planet\tSign\tDegree\tHouse\tMotion'
    main_aspects_header = 'Planet\tAspect\tPlanet\tOrb *\tA/S *'
    other_aspects_header = 'Object\tAspect\tPlanet\tOrb\tAspect'

    # Initialize section indices
    planets_start = -1
    houses_start = -1
    main_aspects_start = -1
    other_aspects_start = -1
    end_index = len(lines)

    # Find the indices of each section
    for idx, line in enumerate(lines):
        if line.strip() == planets_header:
            planets_start = idx
        elif houses_start == -1 and re.match(r'^\d+:', line.strip()):
            houses_start = idx
        elif line.strip() == main_aspects_header:
            main_aspects_start = idx
        elif line.strip() == other_aspects_header:
            other_aspects_start = idx

    # Parse Planets Data
    if planets_start != -1:
        planets_end = houses_start if houses_start != -1 else end_index
        planet_lines = lines[planets_start + 1:planets_end]
        for idx, line in enumerate(planet_lines):
            parts = line.strip().split('\t')
            if len(parts) >= 5:
                planet_line = parts[0].strip(':')
                planet = planet_line.strip()
                sign = parts[1].strip()
                degree = parts[2].strip()
                house = parts[3].strip()
                motion = parts[4].strip()
                # Determine the index as ordinal (first, second, third)
                index = ordinal(idx + 1)
                api_response['planets'].append({
                    'index': index,
                    'planet': planet,
                    'sign': sign,
                    'degree': degree,
                    'house': house,
                    'motion': motion
                })

    # Parse Houses Data
    if houses_start != -1:
        houses_end = main_aspects_start if main_aspects_start != -1 else end_index
        house_lines = lines[houses_start:houses_end]
        for line in house_lines:
            parts = line.strip().split('\t')
            if len(parts) >= 3:
                house_number_line = parts[0].strip(':')
                house_number = house_number_line.strip()
                sign = parts[1].strip()
                degree = parts[2].strip()
                api_response['houses'].append({
                    'index': int(house_number) - 1,
                    'house_number': house_number,
                    'sign': sign,
                    'degree': degree
                })

    # Parse Main Aspects Data
    if main_aspects_start != -1:
        main_aspects_end = other_aspects_start if other_aspects_start != -1 else end_index
        aspect_lines = lines[main_aspects_start + 1:main_aspects_end]
        for idx, line in enumerate(aspect_lines):
            parts = line.strip().split('\t')
            if len(parts) >= 5:
                planet1 = parts[0].strip()
                aspect_type = parts[1].strip()
                planet2 = parts[2].strip()
                orb = parts[3].strip()
                applying_separating = parts[4].strip()
                api_response['main_aspects'].append({
                    'index': idx + 1,  # Start indexing from 1
                    'planet1': planet1,
                    'aspect': aspect_type,
                    'planet2': planet2,
                    'orb': orb,
                    'applying_separating': applying_separating
                })

    # Parse Other Aspects Data
    if other_aspects_start != -1:
        aspect_lines = lines[other_aspects_start + 1:end_index]
        for idx, line in enumerate(aspect_lines):
            parts = line.strip().split('\t')
            if len(parts) >= 5:
                object1 = parts[0].strip()
                aspect_type = parts[1].strip()
                object2 = parts[2].strip()
                orb = parts[3].strip()
                applying_separating = parts[4].strip()
                api_response['other_aspects'].append({
                    'index': idx + 1,  # Start indexing from 1
                    'object1': object1,
                    'aspect': aspect_type,
                    'object2': object2,
                    'orb': orb,
                    'applying_separating': applying_separating
                })

    return {'api_response': api_response, 'user_name': sanitized_user_name}

def ordinal(n):
    """Converts an integer to its ordinal representation."""
    if 10 <= n % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
    return f"{n}{suffix}"

def save_to_jsonl(data, filename='results.jsonl', system_prompt=''):
    """Saves the extracted data to a JSONL file in the specified structure."""
    try:
        with open(filename, 'w', encoding='utf-8') as file:
            total_requests = sum(len(data.get('api_response', {}).get(key, [])) for key in ['planets', 'houses', 'main_aspects', 'other_aspects'])
            custom_id_prefix = 'request-'

            # Extract the relevant data from the API response
            api_response = data.get('api_response')
            if not api_response:
                print("No data available to process.")
                return

            data_items = []

            # Process planets
            planets = api_response.get('planets', [])
            for idx, body in enumerate(planets):
                index = body.get('index')
                planet = body.get('planet')
                sign = body.get('sign')
                degree = body.get('degree')
                house = body.get('house')
                motion = body.get('motion')
                user_content = (f"The {index} planet/body position {planet} is in {sign} at {degree}, "
                                f"in house {house}, with a {motion} motion.")
                data_items.append({'index': idx, 'content': user_content})

            # Process houses
            houses = api_response.get('houses', [])
            for house_info in houses:
                index = house_info.get('index')
                house_number = house_info.get('house_number')
                sign = house_info.get('sign')
                degree = house_info.get('degree')
                user_content = f"House {house_number} ({sign}) is at {degree}."
                data_items.append({'index': index + 1000, 'content': user_content})  # Adding 1000 to index to ensure ordering after planets

            # Process main aspects
            main_aspects = api_response.get('main_aspects', [])
            for aspect in main_aspects:
                index = aspect.get('index')
                planet1 = aspect.get('planet1')
                aspect_type = aspect.get('aspect')
                planet2 = aspect.get('planet2')
                orb = aspect.get('orb')
                applying_separating = aspect.get('applying_separating')
                current_content = (f"{planet1} forms a {aspect_type} with {planet2} "
                                   f"(orb: {orb}, {applying_separating}).")
                user_content = f"Main aspect [{index}] is " + current_content
                data_items.append({'index': index + 2000, 'content': user_content})  # Adding 2000 to index for ordering

            # Process other aspects
            other_aspects = api_response.get('other_aspects', [])
            for aspect in other_aspects:
                index = aspect.get('index')
                object1 = aspect.get('object1')
                aspect_type = aspect.get('aspect')
                object2 = aspect.get('object2')
                orb = aspect.get('orb')
                applying_separating = aspect.get('applying_separating')
                user_content = f"{object1} forms a {aspect_type} with {object2} (orb: {orb}, {applying_separating})."
                data_items.append({'index': index + 3000, 'content': user_content})  # Adding 3000 to index

            if not data_items:
                print("No data available to process.")
                return

            data_items.sort(key=lambda x: x['index'])
            data_length = len(data_items)
            index = 0
            for i in range(1, total_requests + 1):
                # Cycle through the data if there are fewer items than total_requests
                user_content = data_items[index % data_length]['content']
                index += 1

                second_user_content = "Deep dive, motherfucker!"

                messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_content},
                ]

                json_object = {
                    "custom_id": f"{custom_id_prefix}{i}",
                    "method": "POST",
                    "url": "/v1/chat/completions",
                    "body": {
                        "model": "gpt-4o",
                        "messages": messages
                    }
                }

                json_line = json.dumps(json_object, ensure_ascii=True)
                file.write(json_line + '\n')

            print(f"Data saved to {filename} with {total_requests} requests.")
    except Exception as e:
        print(f"An error occurred while saving data to JSONL: {e}")

def extract_assistant_content(json_obj):
    """Extracts the assistant's content from the JSON object."""
    try:
        response_body = json_obj.get('response', {}).get('body', {})
        choices = response_body.get('choices', [])
        if choices:
            assistant_message = choices[0].get('message', {})
            content = assistant_message.get('content', '')
            return content
    except Exception as e:
        print(f"Error extracting assistant content: {e}")
    return None

def clean_content(content):
    """Cleans the assistant's content by handling encoding issues."""
    # Simplify the function by just replacing known unicode characters
    replacements = {
        '\u2019': "'",
        '\u2014': '-',
        '\u201c': '"',
        '\u201d': '"',
        '\u2018': "'",
        '\u2013': '-',
        '\u2026': '...',
        '\u00B0': '°',  # Degree symbol
    }
    # Remove double newlines
    content = content.replace('\n\n', '\n')
    return content

def generate_chart_summary(data):
    """Generates a summary of the chart data."""
    api_response = data.get('api_response', {})
    summary_parts = []

    # Summarize planets
    planets = api_response.get('planets', [])
    for planet_info in planets:
        planet = planet_info.get('planet')
        sign = planet_info.get('sign')
        degree = planet_info.get('degree')
        summary_parts.append(f"{planet} in {sign} at {degree}")

    # Summarize houses
    houses = api_response.get('houses', [])
    for house_info in houses:
        house_number = house_info.get('house_number')
        sign = house_info.get('sign')
        degree = house_info.get('degree')
        summary_parts.append(f"House {house_number}: {sign} at {degree}")

    # Summarize main aspects
    main_aspects = api_response.get('main_aspects', [])
    for aspect_info in main_aspects:
        planet1 = aspect_info.get('planet1')
        aspect = aspect_info.get('aspect')
        planet2 = aspect_info.get('planet2')
        summary_parts.append(f"{planet1} {aspect} {planet2}")

    # Summarize other aspects
    other_aspects = api_response.get('other_aspects', [])
    for aspect_info in other_aspects:
        object1 = aspect_info.get('object1')
        aspect = aspect_info.get('aspect')
        object2 = aspect_info.get('object2')
        summary_parts.append(f"{object1} {aspect} {object2}")

    chart_summary = "; ".join(summary_parts)
    return chart_summary

def get_chat_completion(question, chart_summary, system_prompt):
    """Gets a chat completion from the OpenAI API."""
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    if not client.api_key:
        print("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
        return None

    if not system_prompt:
        system_prompt = "You are a helpful assistant."

    user_content = f"Here's my natal birth chart summary: {chart_summary}. My question for you, Ziggy, is: {question}"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content},
    ]

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=messages
        )
        assistant_message = completion.choices[0].message.content
        return assistant_message
    except Exception as e:
        print(f"Error getting chat completion: {e}")
    return None

def run_generate():
    """Runs the generate command."""
    print("Please paste all your data below. When you're done, press Enter on an empty line to signal the end of input.")
    lines = []
    while True:
        try:
            line = input()
            if line.strip() == '':
                break
            lines.append(line)
        except EOFError:
            break
    all_data = '\n'.join(lines)

    # Parse the input data
    data = parse_input_data(all_data)

    # Get sanitized user name
    sanitized_user_name = data.get('user_name', 'results')

    # Save to JSONL
    filename = f"{sanitized_user_name}.jsonl"
    save_to_jsonl(data, filename=filename, system_prompt=system_prompt)

    # Save the chart summary for later use
    chart_summary = generate_chart_summary(data)
    data['chart_summary'] = chart_summary

    # Save data to a temporary file for later use
    with open(f"{sanitized_user_name}_data.json", 'w', encoding='utf-8') as f:
        json.dump(data, f)

def run_process(input_file, output_file):
    """Runs the process command."""
    # If output_file is not specified, derive it from input_file
    if not output_file:
        filename_without_extension = os.path.splitext(os.path.basename(input_file))[0]
        sanitized_user_name = filename_without_extension
        output_file = f"{sanitized_user_name}.docx"

    # Read the batch JSONL file
    assistant_contents = []

    with open(input_file, 'r', encoding='utf-8') as f:
        for line in f:
            json_obj = json.loads(line)
            # Extract assistant content
            assistant_content = extract_assistant_content(json_obj)
            if assistant_content:
                # Clean up the content
                cleaned_content = clean_content(assistant_content)
                assistant_contents.append(cleaned_content)

    # Write the contents to a docx file
    document = Document()
    for content in assistant_contents:
        document.add_paragraph(content)

    # Save the docx file
    document.save(output_file)
    print(f"Assistant contents saved to {output_file}.")

    # Load chart summary from saved data
    sanitized_user_name = os.path.splitext(os.path.basename(input_file))[0]
    data_file = f"{sanitized_user_name}_data.json"
    if os.path.exists(data_file):
        with open(data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        chart_summary = data.get('chart_summary', '')
    else:
        chart_summary = ''

    # Prompt for personalized questions
    while True:
        add_question = input("Add a personalized question? y/n: ").strip().lower()
        if add_question == 'y':
            question = input("Enter your question: ").strip()
            # Send the question to the OpenAI API, including the chart summary
            assistant_response = get_chat_completion(question, chart_summary, system_prompt)
            if assistant_response:
                # Clean up the content
                cleaned_response = clean_content(assistant_response)
                # Append to the docx file
                document.add_paragraph(cleaned_response)
                document.add_page_break()
                # Save the docx file
                document.save(output_file)
                print("Assistant's response appended to the document.")
        elif add_question == 'n':
            break
        else:
            print("Please enter 'y' or 'n'.")

    print("Processing completed.")

def main():
    parser = argparse.ArgumentParser(description='Process astrology data.')
    subparsers = parser.add_subparsers(dest='command', required=True)

    # Generate command
    generate_parser = subparsers.add_parser('generate', help='Generate JSONL file from input data.')

    # Process command
    process_parser = subparsers.add_parser('process', help='Process batch API response and generate docx file.')
    process_parser.add_argument('--input', help='Path to the batch API JSONL response file.')
    process_parser.add_argument('--output', '-o', help='Output docx file name.')

    args = parser.parse_args()

    if args.command == 'generate':
        run_generate()
    elif args.command == 'process':
        run_process(args.input, args.output)

if __name__ == "__main__":
    main()